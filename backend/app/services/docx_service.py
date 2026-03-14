"""
app/services/docx_service.py
Generazione report .docx lato backend (python-docx)
Chiamato da storage router per creare report on-demand
"""
import logging
from datetime import datetime
from typing import List, Tuple
from io import BytesIO

logger = logging.getLogger(__name__)


def generate_report_docx(
    title: str,
    sections: List[Tuple[str, str]],  # [(heading, content), ...]
    author: str = "Big House AI",
) -> BytesIO:
    """
    Genera un documento .docx in memoria e lo ritorna come BytesIO.

    Args:
        title:    Titolo del documento
        sections: Lista di tuple (intestazione, contenuto)
        author:   Nome autore nel documento

    Returns:
        BytesIO con il file .docx pronto per essere inviato
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Margini ──
        section = doc.sections[0]
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

        # ── Titolo ──
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size  = Pt(22)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6e)  # navy

        # ── Sottotitolo con data ──
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(
            f"Generato da {author} · {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )
        sub_run.font.size  = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        doc.add_paragraph()  # spazio

        # ── Sezioni ──
        for heading, content in sections:
            # Intestazione sezione
            h = doc.add_paragraph()
            h_run = h.add_run(heading)
            h_run.font.size  = Pt(13)
            h_run.font.bold  = True
            h_run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

            # Linea separatrice
            h.paragraph_format.border_bottom = True

            # Contenuto
            para = doc.add_paragraph()
            content_run = para.add_run(content)
            content_run.font.size = Pt(11)

            doc.add_paragraph()  # spazio tra sezioni

        # ── Footer ──
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Big House AI · Documento riservato · {datetime.now().year}"
        )
        footer_run.font.size  = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        # ── Salva in memoria ──
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except ImportError:
        logger.error("python-docx non installato. Aggiungi 'python-docx' a requirements.txt")
        raise RuntimeError("Libreria python-docx non disponibile")
    except Exception as e:
        logger.error(f"Errore generazione DOCX: {e}")
        raise
